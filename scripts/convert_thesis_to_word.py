import os
import re
from pathlib import Path
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def clean_latex_math(math_str):
    """Clean LaTeX math string into elegant mathematical notation for Word."""
    # Common replacements
    math_str = math_str.replace(r'\mathbb{E}', 'E').replace(r'\mathbb{R}', 'ℝ').replace(r'\mathbb{I}', 'I')
    math_str = math_str.replace(r'\boldsymbol{\beta}', 'β').replace(r'\boldsymbol{\gamma}', 'γ').replace(r'\boldsymbol{\eta}', 'η')
    math_str = math_str.replace(r'\boldsymbol{\epsilon}', 'ε').replace(r'\boldsymbol{\theta}', 'θ').replace(r'\boldsymbol{\Omega}', 'Ω')
    math_str = math_str.replace(r'\hat{\boldsymbol{\beta}}', 'β̂').replace(r'\hat{\boldsymbol{\gamma}}', 'γ̂')
    math_str = math_str.replace(r'\alpha', 'α').replace(r'\beta', 'β').replace(r'\gamma', 'γ').replace(r'\delta', 'δ')
    math_str = math_str.replace(r'\lambda', 'λ').replace(r'\sigma', 'σ').replace(r'\tau', 'τ').replace(r'\omega', 'ω')
    math_str = math_str.replace(r'\hat{\alpha}', 'α̂').replace(r'\hat{\beta}', 'β̂').replace(r'\hat{\gamma}', 'γ̂')
    math_str = math_str.replace(r'\sum_{i=1}^{N_t}', '∑ (i=1..N)').replace(r'\sum_{k=1}^{K}', '∑ (k=1..K)')
    math_str = math_str.replace(r'\sum', '∑').replace(r'\prod', '∏').replace(r'\int', '∫')
    math_str = math_str.replace(r'\arg\min', 'argmin').replace(r'\arg\max', 'argmax')
    math_str = math_str.replace(r'\operatorname{Softmax}', 'Softmax').replace(r'\operatorname{Huber}', 'Huber')
    math_str = math_str.replace(r'\operatorname{Var}', 'Var').replace(r'\operatorname{Cov}', 'Cov')
    math_str = math_str.replace(r'\operatorname{Rank}', 'Rank').replace(r'\operatorname{sign}', 'sign')
    math_str = math_str.replace(r'\operatorname{diag}', 'diag').replace(r'\operatorname{Gain}', 'Gain')
    math_str = math_str.replace(r'\mathbf{', '').replace(r'\text{', '').replace(r'\mathrm{', '').replace(r'\mathcal{', '')
    math_str = math_str.replace(r'\}', '').replace(r'\{', '')
    math_str = math_str.replace(r'\in', '∈').replace(r'\ge', '≥').replace(r'\le', '≤').replace(r'\to', '→')
    math_str = math_str.replace(r'\cdot', '·').replace(r'\odot', '⊙').replace(r'\otimes', '⊗')
    math_str = math_str.replace(r'\infty', '∞').replace(r'\approx', '≈').replace(r'\neq', '≠')
    math_str = math_str.replace(r'\quad', '   ').replace(r'\qquad', '      ').replace(r'\_', '_')
    math_str = re.sub(r'\\label\{[^}]+\}', '', math_str)
    return math_str.strip()

def clean_latex_text(text):
    text = re.sub(r'%.*', '', text)
    text = text.replace(r'\textbf{', '').replace(r'\textit{', '').replace(r'\texttt{', '').replace(r'\path{', '')
    text = text.replace(r'\allowbreak', '').replace(r'\_', '_').replace(r'\&', '&').replace(r'\%', '%')
    text = text.replace(r'\$', '$').replace(r'\}', '}').replace(r'\{', '{')
    text = re.sub(r'\\cite[t|p]?\*?\{([^}]+)\}', r'(\1)', text)
    text = re.sub(r'\\ref\{([^}]+)\}', r'\1', text)
    text = re.sub(r'\\label\{([^}]+)\}', '', text)
    text = re.sub(r'\\begin\{[^}]+\}', '', text)
    text = re.sub(r'\\end\{[^}]+\}', '', text)
    text = re.sub(r'\\item', '• ', text)
    return text.strip()

def build_word_thesis():
    root_dir = Path("c:/Users/murta/Desktop/thesis final 2.0")
    thesis_dir = root_dir / "thesis"
    output_docx = root_dir / "thesis.docx"
    
    doc = docx.Document()
    
    # Margins
    for s in doc.sections:
        s.top_margin = Inches(1.25)
        s.bottom_margin = Inches(1.25)
        s.left_margin = Inches(1.25)
        s.right_margin = Inches(1.25)
        
    COLOR_NAVY = RGBColor(10, 34, 64)
    COLOR_GOLD = RGBColor(175, 140, 65)
    COLOR_SLATE = RGBColor(70, 90, 110)
    COLOR_DARK = RGBColor(35, 40, 45)
    
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11.5)
    style_normal.font.color.rgb = COLOR_DARK
    style_normal.paragraph_format.line_spacing = 1.3
    style_normal.paragraph_format.space_after = Pt(6)
    
    # ── 1. TITLE PAGE ──────────────────────────────────────────────────────────
    p_uni = doc.add_paragraph()
    p_uni.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_uni = p_uni.add_run("UNIVERSITY OF VERONA\n")
    r_uni.bold = True
    r_uni.font.size = Pt(16)
    r_uni.font.color.rgb = COLOR_NAVY
    
    r_dept = p_uni.add_run("DEPARTMENT OF ECONOMICS\n")
    r_dept.font.size = Pt(13)
    r_dept.font.color.rgb = COLOR_SLATE
    
    r_deg = p_uni.add_run("Master's Degree in Economics and Data Analysis\n\n")
    r_deg.font.size = Pt(11)
    r_deg.font.color.rgb = COLOR_SLATE
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(18)
    
    r_t1 = p_title.add_run("DO MACHINE LEARNING MODELS IMPROVE STOCK RETURN PREDICTION?\n")
    r_t1.bold = True
    r_t1.font.size = Pt(20)
    r_t1.font.color.rgb = COLOR_NAVY
    
    r_t2 = p_title.add_run("Evidence from S&P 500 Constituent Markets and Dimension Sensitivity (2015–2024)\n")
    r_t2.font.size = Pt(13)
    r_t2.font.italic = True
    r_t2.font.color.rgb = COLOR_SLATE
    
    doc.add_paragraph()
    meta_table = doc.add_table(rows=1, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_cand, cell_sup = meta_table.rows[0].cells
    cell_cand.width = Inches(3.0)
    cell_sup.width = Inches(3.0)
    
    p_c = cell_cand.paragraphs[0]
    r = p_c.add_run("Candidate:\n")
    r.bold = True
    r.font.color.rgb = COLOR_NAVY
    p_c.add_run("Murtuza Yusuf Rangwala\nMatricola: VR508566")
    
    p_s = cell_sup.paragraphs[0]
    p_s.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p_s.add_run("Supervisor:\n")
    r.bold = True
    r.font.color.rgb = COLOR_NAVY
    p_s.add_run("Prof. Giuseppina Chesini\nDepartment of Economics\nUniversity of Verona")
    
    doc.add_paragraph()
    p_yr = doc.add_paragraph()
    p_yr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_yr = p_yr.add_run("Academic Year 2024/2025")
    r_yr.font.size = Pt(10)
    r_yr.font.color.rgb = COLOR_SLATE
    
    doc.add_page_break()
    
    # ── 2. ABSTRACT ────────────────────────────────────────────────────────────
    h_abs = doc.add_heading("Abstract", level=1)
    h_abs.runs[0].font.color.rgb = COLOR_NAVY
    h_abs.runs[0].font.size = Pt(18)
    
    p_abs = doc.add_paragraph()
    p_abs.paragraph_format.line_spacing = 1.3
    p_abs.add_run(
        "This thesis investigates whether non-linear machine learning models and deep sequence architectures enhance "
        "cross-sectional stock return prediction relative to traditional linear factor pricing models. Utilizing a "
        "point-in-time aligned dataset of S&P 500 constituent equities spanning 2015 through 2024, we construct a 53-variable "
        "multimodal feature space combining technical indicators, fundamental accounting ratios, sentiment scores, and macroeconomic variables. "
        "To evaluate out-of-sample forecast accuracy without backtest contamination, models are estimated using a 5-fold expanding-window "
        "walk-forward validation scheme.\n\n"
        "We introduce a novel deep learning architecture—the Temporal Fusion Deep Multimodal Gated Attention (TFDMGA) network—featuring "
        "Causal TCN encoders, a sequential directed ring attention cascade, and a 3-way macro-conditioned dynamic gating module. "
        "Out-of-sample empirical results demonstrate that deep sequence models achieve statistically significant predictive superiority "
        "over linear baselines (Fama-MacBeth OLS, LASSO) and decision tree ensembles (Random Forest, XGBoost), yielding higher "
        "Information Coefficients and directional accuracy. Under institutional transaction cost friction (10 bps) and a 1-day execution "
        "buffer, integrating a 2:1 Take-Profit/Stop-Loss risk management overlay preserves capital during major macroeconomic drawdowns.\n\n"
        "Finally, Fama-French 5-factor spanning regressions reveal that strategy returns yield a statistically zero net alpha (α̂ = -0.18% per annum, "
        "p = 0.976, R² = 41.2%), with 100% of return variation accounted for by dynamic systematic factor risk exposures (specifically operating profitability RMW, t = +9.12). "
        "These findings demonstrate that machine learning models function as dynamic factor timing engines rather than discoverers of unpriced market arbitrage, "
        "maintaining consistency with market efficiency."
    )
    
    p_kw = doc.add_paragraph()
    p_kw.paragraph_format.space_before = Pt(12)
    r_kw = p_kw.add_run("Keywords: ")
    r_kw.bold = True
    p_kw.add_run("Machine Learning, Empirical Asset Pricing, Cross-Sectional Return Prediction, Deep Learning, Transformer Attention, Fama-French Spanning, Transaction Costs, Walk-Forward Optimization.")
    
    doc.add_page_break()
    
    # ── 3. PROCESS CHAPTER FILES ───────────────────────────────────────────────
    chapters = [
        ("introduction.tex", "Chapter 1: Introduction"),
        ("literature_review.tex", "Chapter 2: Literature Review"),
        ("data_and_features.tex", "Chapter 3: Data & Feature Engineering Architecture"),
        ("methodology.tex", "Chapter 4: Methodology & Model Architectures"),
        ("baseline_results.tex", "Chapter 5: Baseline Econometric Results & Feature Screening"),
        ("ml_and_backtesting.tex", "Chapter 6: Out-of-Sample Performance, Backtesting & Portfolio Compounding"),
        ("conclusion.tex", "Chapter 7: Conclusion & Doctoral Research Roadmap")
    ]
    
    eq_counter = 1
    
    for filename, chap_title in chapters:
        filepath = thesis_dir / "chapters" / filename
        if not filepath.exists():
            continue
            
        with open(filepath, "r", encoding="utf-8") as f:
            lines = f.readlines()
            
        print(f"Formatting equations & text for {filename}...")
        
        h_ch = doc.add_heading(chap_title, level=1)
        h_ch.runs[0].font.color.rgb = COLOR_NAVY
        h_ch.runs[0].font.size = Pt(18)
        h_ch.paragraph_format.space_before = Pt(18)
        h_ch.paragraph_format.space_after = Pt(12)
        
        in_equation = False
        in_figure = False
        eq_buffer = []
        fig_caption = ""
        fig_image = ""
        
        for line in lines:
            line_str = line.strip()
            
            if line_str.startswith("%") or line_str.startswith(r"\chapter"):
                continue
                
            if line_str.startswith(r"\section{"):
                sec_title = re.findall(r"\\section\{([^}]+)\}", line_str)
                if sec_title:
                    h2 = doc.add_heading(clean_latex_text(sec_title[0]), level=2)
                    h2.runs[0].font.color.rgb = COLOR_NAVY
                    h2.runs[0].font.size = Pt(14)
                    h2.paragraph_format.space_before = Pt(14)
                    h2.paragraph_format.space_after = Pt(6)
                continue
                
            if line_str.startswith(r"\subsection{"):
                subsec_title = re.findall(r"\\subsection\{([^}]+)\}", line_str)
                if subsec_title:
                    h3 = doc.add_heading(clean_latex_text(subsec_title[0]), level=3)
                    h3.runs[0].font.color.rgb = COLOR_SLATE
                    h3.runs[0].font.size = Pt(12)
                    h3.paragraph_format.space_before = Pt(10)
                    h3.paragraph_format.space_after = Pt(4)
                continue

            # Figures
            if r"\begin{figure}" in line_str:
                in_figure = True
                fig_caption = ""
                fig_image = ""
                continue
            if in_figure:
                if r"\includegraphics" in line_str:
                    img_match = re.search(r"\\includegraphics\[[^\]]*\]\{([^}]+)\}", line_str)
                    if img_match:
                        fig_image = img_match.group(1).replace("./figures/", "").replace("figures/", "")
                if r"\caption{" in line_str:
                    cap_match = re.search(r"\\caption\{([^}]+)\}", line_str)
                    if cap_match:
                        fig_caption = clean_latex_text(cap_match.group(1))
                if r"\end{figure}" in line_str:
                    in_figure = False
                    if fig_image:
                        img_path = thesis_dir / "figures" / fig_image
                        if img_path.exists():
                            p_img = doc.add_paragraph()
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_img.paragraph_format.space_before = Pt(12)
                            p_img.paragraph_format.space_after = Pt(4)
                            try:
                                p_img.add_run().add_picture(str(img_path), width=Inches(5.5))
                            except Exception:
                                pass
                    if fig_caption:
                        p_cap = doc.add_paragraph()
                        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r_cap = p_cap.add_run(f"Figure: {fig_caption}")
                        r_cap.font.italic = True
                        r_cap.font.size = Pt(10)
                        r_cap.font.color.rgb = COLOR_SLATE
                        p_cap.paragraph_format.space_after = Pt(12)
                continue

            # ── ENHANCED EQUATION FORMATTING IN WORD ──────────────────────────────
            if r"\begin{equation}" in line_str or r"\begin{align}" in line_str or r"\begin{split}" in line_str:
                in_equation = True
                eq_buffer = []
                continue
            if in_equation:
                if r"\end{equation}" in line_str or r"\end{align}" in line_str or r"\end{split}" in line_str:
                    in_equation = False
                    if eq_buffer:
                        full_eq_text = " ".join(eq_buffer)
                        cleaned_eq = clean_latex_math(full_eq_text)
                        
                        # Create a dedicated math box table in Word for high-grade equation styling
                        eq_table = doc.add_table(rows=1, cols=2)
                        eq_table.alignment = WD_TABLE_ALIGNMENT.CENTER
                        eq_table.autofit = False
                        
                        c_eq, c_num = eq_table.rows[0].cells
                        c_eq.width = Inches(5.5)
                        c_num.width = Inches(1.0)
                        
                        set_cell_background(c_eq, "F5F7FA")
                        set_cell_background(c_num, "F5F7FA")
                        set_cell_margins(c_eq, top=80, bottom=80, left=120, right=120)
                        set_cell_margins(c_num, top=80, bottom=80, left=120, right=120)
                        
                        p_eq_cell = c_eq.paragraphs[0]
                        p_eq_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r_m = p_eq_cell.add_run(cleaned_eq)
                        r_m.font.name = 'Cambria Math'
                        r_m.font.size = Pt(11)
                        r_m.font.italic = True
                        r_m.font.color.rgb = COLOR_NAVY
                        
                        p_num_cell = c_num.paragraphs[0]
                        p_num_cell.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        r_n = p_num_cell.add_run(f"({eq_counter})")
                        r_n.font.name = 'Calibri'
                        r_n.font.size = Pt(10)
                        r_n.font.bold = True
                        r_n.font.color.rgb = COLOR_SLATE
                        
                        eq_counter += 1
                        doc.add_paragraph() # spacing
                else:
                    eq_buffer.append(line_str)
                continue

            # Standard Text Paragraphs
            cleaned_text = clean_latex_text(line_str)
            if cleaned_text and not cleaned_text.startswith(r"\label") and not cleaned_text.startswith(r"\begin") and not cleaned_text.startswith(r"\end"):
                p_text = doc.add_paragraph()
                p_text.paragraph_format.line_spacing = 1.3
                p_text.paragraph_format.space_after = Pt(6)
                if cleaned_text.startswith("• "):
                    p_text.paragraph_format.left_indent = Inches(0.25)
                    p_text.add_run(cleaned_text[2:])
                else:
                    p_text.add_run(cleaned_text)

        doc.add_page_break()

    # Bibliography
    h_bib = doc.add_heading("References", level=1)
    h_bib.runs[0].font.color.rgb = COLOR_NAVY
    h_bib.runs[0].font.size = Pt(18)
    
    bib_path = thesis_dir / "references.bib"
    if bib_path.exists():
        with open(bib_path, "r", encoding="utf-8") as f:
            bib_content = f.read()
            
        entries = re.findall(r"@\w+\{([^,]+),\s*author\s*=\s*\{([^}]+)\},\s*title\s*=\s*\{([^}]+)\},\s*journal\s*=\s*\{([^}]+)\},\s*year\s*=\s*\{([^}]+)\}", bib_content)
        for key, author, title, journal, year in entries[:50]:
            p_b = doc.add_paragraph()
            p_b.paragraph_format.left_indent = Inches(0.4)
            p_b.paragraph_format.first_line_indent = Inches(-0.4)
            p_b.paragraph_format.space_after = Pt(4)
            r_auth = p_b.add_run(f"{clean_latex_text(author)} ({year}). ")
            r_auth.bold = True
            p_b.add_run(f"{clean_latex_text(title)}. ")
            r_j = p_b.add_run(f"{clean_latex_text(journal)}.")
            r_j.font.italic = True

    doc.save(output_docx)
    print(f"Successfully generated Master Thesis Word Document with formatted math: {output_docx}")

if __name__ == "__main__":
    build_word_thesis()
